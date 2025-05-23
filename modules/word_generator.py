import os
from datetime import datetime
from docx import Document
from pathlib import Path
from config import TEMPLATE_OBRA_SELLADO, TEMPLATE_OBRA_VISADO, TEMPLATE_INFORME

class WordGenerator:
    def __init__(self):
        self.template_obra_sellado = TEMPLATE_OBRA_SELLADO
        self.template_obra_visado = TEMPLATE_OBRA_VISADO
        self.template_informe = TEMPLATE_INFORME
    
    def _check_templates(self):
        """Verifica que las plantillas existan"""
        missing = []
        if not os.path.exists(self.template_obra_sellado):
            missing.append("Plantilla de obra sellado")
        if not os.path.exists(self.template_obra_visado):
            missing.append("Plantilla de obra visado")
        if not os.path.exists(self.template_informe):
            missing.append("Plantilla de informe técnico")
        
        return missing
    
    def generate_obra_sellado(self, data, output_path=None):
        """
        Genera un documento Word para la tasa de sellado de una obra manteniendo el formato.
        
        Args:
            data: Diccionario con los datos de la obra
            output_path: Ruta donde guardar el documento (opcional)
        
        Returns:
            str: Ruta al documento generado
        """
        missing = self._check_templates()
        if missing:
            return f"Faltan plantillas: {', '.join(missing)}"
        
        try:
            # Cargar la plantilla
            doc = Document(self.template_obra_sellado)
            
            # Preparar los reemplazos
            replacements = {
                "[FECHA]": str(data.get("fecha", "") or ""),
                "[PROFESION]": str(data.get("profesion", "") or ""),
                "[NOMBRE_PROFESIONAL]": str(data.get("nombre_profesional", "") or ""),
                "[NOMBRE_COMITENTE]": str(data.get("nombre_comitente", "") or ""),
                "[UBICACION]": str(data.get("ubicacion", "") or ""),
                "[NRO_EXPTE_MUNICIPAL]": str(data.get("nro_expte_municipal", "") or ""),
                "[NRO_SISTEMA_GOP]": str(data.get("nro_sistema_gop", "") or ""),
                "[NRO_PARTIDA_INMOBILIARIA]": str(data.get("nro_partida_inmobiliaria", "") or ""),
                "[TASA_SELLADO]": str(data.get("tasa_sellado", "") or "")
            }
            
            # Método para reemplazar en un párrafo respetando el formato por run
            def replace_in_paragraph_by_runs(paragraph):
                # Para cada run, reemplazamos los marcadores conservando su formato
                modified = False
                for run in paragraph.runs:
                    original_text = run.text
                    modified_text = original_text
                    
                    for placeholder, value in replacements.items():
                        if placeholder in modified_text:
                            modified_text = modified_text.replace(placeholder, value)
                            modified = True
                    
                    # Solo actualizar si hubo cambios
                    if modified_text != original_text:
                        run.text = modified_text
                
                return modified
            
            # Método alternativo en caso de problemas con marcadores específicos
            def clean_and_replace_paragraph(paragraph):
                original_text = paragraph.text
                modified_text = original_text
                
                # Eliminar cualquier duplicación existente que pudo ocurrir en ediciones anteriores
                for placeholder, value in replacements.items():
                    # Buscar patrones como "[FECHA]25/04/2025"
                    pattern = f"{placeholder}{value}"
                    if pattern in modified_text:
                        modified_text = modified_text.replace(pattern, value)
                
                # Hacer los reemplazos normales
                for placeholder, value in replacements.items():
                    modified_text = modified_text.replace(placeholder, value)
                
                # Si el texto cambió, actualizamos el párrafo
                if modified_text != original_text:
                    # Guardar formatos
                    formats = []
                    for run in paragraph.runs:
                        formats.append({
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font': run.font.name,
                            'size': run.font.size,
                            'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                        })
                    
                    # Calcular formato predominante
                    predominant_format = {
                        'bold': sum(1 for f in formats if f['bold'] == True) > len(formats) / 2,
                        'italic': sum(1 for f in formats if f['italic'] == True) > len(formats) / 2,
                        'underline': sum(1 for f in formats if f['underline'] == True) > len(formats) / 2
                    }
                    
                    # Si hay algún formato, elegir el del primer run como base
                    if formats:
                        predominant_format['font'] = formats[0]['font']
                        predominant_format['size'] = formats[0]['size']
                        predominant_format['color'] = formats[0]['color']
                    
                    # Limpiar todos los runs
                    while len(paragraph.runs) > 0:
                        paragraph._p.remove(paragraph.runs[0]._r)
                    
                    # Agregar nuevo run con el texto modificado
                    run = paragraph.add_run(modified_text)
                    
                    # Aplicar formato
                    run.bold = predominant_format.get('bold', False)
                    run.italic = predominant_format.get('italic', False)
                    run.underline = predominant_format.get('underline', False)
                    if predominant_format.get('font'):
                        run.font.name = predominant_format['font']
                    if predominant_format.get('size'):
                        run.font.size = predominant_format['size']
                    if predominant_format.get('color'):
                        run.font.color.rgb = predominant_format['color']
                    
                    return True
                
                return False
            
            # Procesar párrafos normales
            for paragraph in doc.paragraphs:
                # Buscar patrones de duplicación
                contains_duplicates = False
                for placeholder, value in replacements.items():
                    pattern = f"{placeholder}{value}"
                    if pattern in paragraph.text:
                        contains_duplicates = True
                        break
                
                if "[FECHA]" in paragraph.text or contains_duplicates:
                    # Usar el enfoque limpio para FECHA y párrafos con duplicación
                    clean_and_replace_paragraph(paragraph)
                else:
                    # Usar el enfoque por runs para el resto
                    replace_in_paragraph_by_runs(paragraph)
            
            # Procesar tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Buscar patrones de duplicación
                            contains_duplicates = False
                            for placeholder, value in replacements.items():
                                pattern = f"{placeholder}{value}"
                                if pattern in paragraph.text:
                                    contains_duplicates = True
                                    break
                            
                            if "[FECHA]" in paragraph.text or contains_duplicates:
                                # Usar el enfoque limpio para FECHA y párrafos con duplicación
                                clean_and_replace_paragraph(paragraph)
                            else:
                                # Usar el enfoque por runs para el resto
                                replace_in_paragraph_by_runs(paragraph)
            
            # Determinar la ruta de salida
            if output_path is None:
                output_dir = Path.cwd()
                output_path = output_dir / f"Sellado_{data.get('nombre_comitente', 'Comitente')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # Guardar el documento
            doc.save(output_path)
            return str(output_path)
        
        except Exception as e:
            return f"Error al generar el documento: {str(e)}"
    
    def generate_obra_visado(self, data, output_path=None):
        """
        Genera un documento Word para la tasa de visado de una obra manteniendo el formato.
        
        Args:
            data: Diccionario con los datos de la obra
            output_path: Ruta donde guardar el documento (opcional)
        
        Returns:
            str: Ruta al documento generado
        """
        missing = self._check_templates()
        if missing:
            return f"Faltan plantillas: {', '.join(missing)}"
        
        try:
            # Cargar la plantilla
            doc = Document(self.template_obra_visado)
            
            # Preparar los reemplazos
            replacements = {
                "[FECHA]": str(data.get("fecha", "") or ""),
                "[PROFESION]": str(data.get("profesion", "") or ""),
                "[NOMBRE_PROFESIONAL]": str(data.get("nombre_profesional", "") or ""),
                "[NOMBRE_COMITENTE]": str(data.get("nombre_comitente", "") or ""),
                "[UBICACION]": str(data.get("ubicacion", "") or ""),
                "[NRO_EXPTE_MUNICIPAL]": str(data.get("nro_expte_municipal", "") or ""),
                "[NRO_SISTEMA_GOP]": str(data.get("nro_sistema_gop", "") or ""),
                "[NRO_PARTIDA_INMOBILIARIA]": str(data.get("nro_partida_inmobiliaria", "") or ""),
                "[TASA_VISADO]": str(data.get("tasa_visado", "") or "")
            }
            
            # Método para reemplazar en un párrafo respetando el formato por run
            def replace_in_paragraph_by_runs(paragraph):
                # Para cada run, reemplazamos los marcadores conservando su formato
                modified = False
                for run in paragraph.runs:
                    original_text = run.text
                    modified_text = original_text
                    
                    for placeholder, value in replacements.items():
                        if placeholder in modified_text:
                            modified_text = modified_text.replace(placeholder, value)
                            modified = True
                    
                    # Solo actualizar si hubo cambios
                    if modified_text != original_text:
                        run.text = modified_text
                
                return modified
            
            # Método alternativo en caso de problemas con marcadores específicos
            def clean_and_replace_paragraph(paragraph):
                original_text = paragraph.text
                modified_text = original_text
                
                # Eliminar cualquier duplicación existente que pudo ocurrir en ediciones anteriores
                for placeholder, value in replacements.items():
                    # Buscar patrones como "[FECHA]25/04/2025"
                    pattern = f"{placeholder}{value}"
                    if pattern in modified_text:
                        modified_text = modified_text.replace(pattern, value)
                
                # Hacer los reemplazos normales
                for placeholder, value in replacements.items():
                    modified_text = modified_text.replace(placeholder, value)
                
                # Si el texto cambió, actualizamos el párrafo
                if modified_text != original_text:
                    # Guardar formatos
                    formats = []
                    for run in paragraph.runs:
                        formats.append({
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font': run.font.name,
                            'size': run.font.size,
                            'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                        })
                    
                    # Calcular formato predominante
                    predominant_format = {
                        'bold': sum(1 for f in formats if f['bold'] == True) > len(formats) / 2,
                        'italic': sum(1 for f in formats if f['italic'] == True) > len(formats) / 2,
                        'underline': sum(1 for f in formats if f['underline'] == True) > len(formats) / 2
                    }
                    
                    # Si hay algún formato, elegir el del primer run como base
                    if formats:
                        predominant_format['font'] = formats[0]['font']
                        predominant_format['size'] = formats[0]['size']
                        predominant_format['color'] = formats[0]['color']
                    
                    # Limpiar todos los runs
                    while len(paragraph.runs) > 0:
                        paragraph._p.remove(paragraph.runs[0]._r)
                    
                    # Agregar nuevo run con el texto modificado
                    run = paragraph.add_run(modified_text)
                    
                    # Aplicar formato
                    run.bold = predominant_format.get('bold', False)
                    run.italic = predominant_format.get('italic', False)
                    run.underline = predominant_format.get('underline', False)
                    if predominant_format.get('font'):
                        run.font.name = predominant_format['font']
                    if predominant_format.get('size'):
                        run.font.size = predominant_format['size']
                    if predominant_format.get('color'):
                        run.font.color.rgb = predominant_format['color']
                    
                    return True
                
                return False
            
            # Procesar párrafos normales
            for paragraph in doc.paragraphs:
                # Buscar patrones de duplicación
                contains_duplicates = False
                for placeholder, value in replacements.items():
                    pattern = f"{placeholder}{value}"
                    if pattern in paragraph.text:
                        contains_duplicates = True
                        break
                
                if "[FECHA]" in paragraph.text or contains_duplicates:
                    # Usar el enfoque limpio para FECHA y párrafos con duplicación
                    clean_and_replace_paragraph(paragraph)
                else:
                    # Usar el enfoque por runs para el resto
                    replace_in_paragraph_by_runs(paragraph)
            
            # Procesar tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Buscar patrones de duplicación
                            contains_duplicates = False
                            for placeholder, value in replacements.items():
                                pattern = f"{placeholder}{value}"
                                if pattern in paragraph.text:
                                    contains_duplicates = True
                                    break
                            
                            if "[FECHA]" in paragraph.text or contains_duplicates:
                                # Usar el enfoque limpio para FECHA y párrafos con duplicación
                                clean_and_replace_paragraph(paragraph)
                            else:
                                # Usar el enfoque por runs para el resto
                                replace_in_paragraph_by_runs(paragraph)
            
            # Determinar la ruta de salida
            if output_path is None:
                output_dir = Path.cwd()
                output_path = output_dir / f"Visado_{data.get('nombre_comitente', 'Comitente')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # Guardar el documento
            doc.save(output_path)
            return str(output_path)
        
        except Exception as e:
            return f"Error al generar el documento: {str(e)}"
        
    
    
    def generate_informe(self, data, output_path=None):
        """
        Genera un documento Word para un informe técnico manteniendo el formato.
        
        Args:
            data: Diccionario con los datos del informe
            output_path: Ruta donde guardar el documento (opcional)
        
        Returns:
            str: Ruta al documento generado
        """
        missing = self._check_templates()
        if missing:
            return f"Faltan plantillas: {', '.join(missing)}"
        
        try:
            # Cargar la plantilla
            doc = Document(self.template_informe)
            
            # Función para obtener el valor de reemplazo
            def get_replacement(placeholder):
                if placeholder == "[FECHA]":
                    return str(data.get("fecha", "") or "")
                elif placeholder == "[PROFESION]":
                    return str(data.get("profesion", "") or "")
                elif placeholder == "[PROFESIONAL]":
                    return str(data.get("profesional", "") or "")
                elif placeholder == "[COMITENTE]":
                    return str(data.get("comitente", "") or "")
                elif placeholder == "[TIPO_TRABAJO]":
                    return str(data.get("tipo_trabajo", "") or "")
                elif placeholder == "[DETALLE]":
                    return str(data.get("detalle", "") or "")
                elif placeholder == "[TASA_SELLADO]":
                    return str(data.get("tasa_sellado", "") or "")
                return placeholder  # Si no coincide con ninguno, devuelve el original
            
            # Función para reemplazar marcadores manteniendo el formato
            def replace_in_paragraph(paragraph):
                # Buscar patrones como [ALGO] en el texto
                import re
                placeholder_pattern = r'\[([^\]]+)\]'
                
                # Buscar placeholders en el párrafo
                paragraph_text = paragraph.text
                matches = list(re.finditer(placeholder_pattern, paragraph_text))
                
                # Si no hay coincidencias, no hacer nada
                if not matches:
                    return
                
                # Reemplazar desde el final para no afectar los índices
                for match in reversed(matches):
                    placeholder = match.group(0)  # El texto completo [ALGO]
                    start_idx = match.start()
                    end_idx = match.end()
                    
                    # Encontrar el run que contiene el placeholder
                    current_idx = 0
                    start_run_idx = None
                    end_run_idx = None
                    placeholder_runs = []
                    
                    # Encontrar los runs que contienen el placeholder
                    for i, run in enumerate(paragraph.runs):
                        run_len = len(run.text)
                        
                        # Si el placeholder comienza en este run
                        if start_run_idx is None and start_idx >= current_idx and start_idx < current_idx + run_len:
                            start_run_idx = i
                            start_in_run = start_idx - current_idx
                        
                        # Si el placeholder termina en este run
                        if start_run_idx is not None and end_run_idx is None and end_idx <= current_idx + run_len:
                            end_run_idx = i
                            end_in_run = end_idx - current_idx
                            break
                        
                        current_idx += run_len
                    
                    # Si no encontramos dónde está el placeholder, continuar
                    if start_run_idx is None or end_run_idx is None:
                        continue
                    
                    # Obtener el valor de reemplazo
                    replacement = get_replacement(placeholder)
                    
                    # Caso especial: placeholder contenido en un solo run
                    if start_run_idx == end_run_idx:
                        run = paragraph.runs[start_run_idx]
                        run.text = run.text[:start_in_run] + replacement + run.text[end_in_run:]
                    else:
                        # Caso multiple run: manejar el primer run
                        run = paragraph.runs[start_run_idx]
                        run.text = run.text[:start_in_run] + replacement
                        
                        # Manejar runs intermedios
                        for i in range(start_run_idx + 1, end_run_idx):
                            paragraph.runs[i].text = ""
                        
                        # Manejar el último run
                        if end_run_idx > start_run_idx:
                            run = paragraph.runs[end_run_idx]
                            run.text = run.text[end_in_run:]
            
            # Reemplazar en párrafos normales
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph)
            
            # Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)
            
            # Determinar la ruta de salida
            if output_path is None:
                output_dir = Path.cwd()
                output_path = output_dir / f"Informe_{data.get('comitente', 'Comitente')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # Guardar el documento
            doc.save(output_path)
            return str(output_path)
        
        except Exception as e:
            return f"Error al generar el documento: {str(e)}"